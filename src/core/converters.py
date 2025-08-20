"""Format converters for Markdown Exporter."""

from abc import ABC, abstractmethod
from pathlib import Path
from typing import Optional
import re
import os
import tempfile
import subprocess
import logging

# Import header/footer configuration
try:
    from src.config.header_footer_config import HeaderFooterConfig
    HEADER_FOOTER_AVAILABLE = True
except ImportError:
    HEADER_FOOTER_AVAILABLE = False

# Import conversion libraries
try:
    import docx
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    CONVERSION_AVAILABLE = True
    logger = logging.getLogger(__name__)
    logger.info("✅ Conversion libraries imported successfully")
except ImportError as e:
    logger = logging.getLogger(__name__)
    logger.error(f"❌ Conversion libraries not available: {e}")
    CONVERSION_AVAILABLE = False


def render_mermaid_diagram(mermaid_code: str, output_dir: str) -> str:
    """Render Mermaid diagram to PNG image."""
    try:
        # Create temporary file for Mermaid code
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False, encoding='utf-8') as f:
            f.write(mermaid_code)
            mermaid_file = f.name
        
        # Generate unique filename for output
        import uuid
        output_file = os.path.join(output_dir, f"diagram_{uuid.uuid4().hex[:8]}.png")
        
        # Use mmdc to render the diagram
        # Try to find mmdc in common locations
        mmdc_paths = [
            'mmdc',  # Try system PATH first
            'C:\\Users\\sovan\\AppData\\Roaming\\npm\\mmdc.cmd',  # Global npm (Windows)
            'C:\\Users\\sovan\\AppData\\Roaming\\npm\\mmdc.exe',  # Global npm (Windows exe)
            os.path.join('.venv', 'Scripts', 'mmdc'),  # Virtual env
        ]
        
        mmdc_path = 'mmdc'  # Default to system PATH
        for path in mmdc_paths:
            try:
                result = subprocess.run([path, '--version'], capture_output=True, text=True)
                if result.returncode == 0:
                    mmdc_path = path
                    logger.info(f"Found mmdc at: {mmdc_path}")
                    break
            except Exception as e:
                logger.debug(f"mmdc not found at {path}: {e}")
                continue
        
        cmd = [
            mmdc_path,
            '-i', mermaid_file,
            '-o', output_file,
            '-b', 'transparent'
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True)
        
        # Clean up temporary file
        os.unlink(mermaid_file)
        
        if result.returncode == 0 and os.path.exists(output_file):
            return output_file
        else:
            logger.warning(f"Failed to render Mermaid diagram: {result.stderr}")
            return None
            
    except Exception as e:
        logger.error(f"Error rendering Mermaid diagram: {e}")
        return None


def process_formatted_text(text, format_type="html"):
    """Process text with proper formatting (bold, italic, code, hyperlinks)."""
    if format_type == "html":
        # For PDF and HTML output
        result_parts = []
        current_pos = 0
        
        while current_pos < len(text):
            # Find the next formatting marker
            bold_match = re.search(r'\*\*(.*?)\*\*', text[current_pos:])
            italic_match = re.search(r'\*(.*?)\*', text[current_pos:])
            code_match = re.search(r'`(.*?)`', text[current_pos:])
            link_match = re.search(r'\[(.*?)\]\((.*?)\)', text[current_pos:])
            
            # Find the earliest match
            matches = []
            if bold_match:
                matches.append(('bold', bold_match.start() + current_pos, bold_match))
            if italic_match:
                matches.append(('italic', italic_match.start() + current_pos, italic_match))
            if code_match:
                matches.append(('code', code_match.start() + current_pos, code_match))
            if link_match:
                matches.append(('link', link_match.start() + current_pos, link_match))
            
            if not matches:
                # No more formatting, add remaining text
                if current_pos < len(text):
                    result_parts.append(text[current_pos:])
                break
            
            # Sort by position and take the earliest
            matches.sort(key=lambda x: x[1])
            format_type, pos, match = matches[0]
            
            # Add text before the formatting
            if pos > current_pos:
                result_parts.append(text[current_pos:pos])
            
            # Add formatted text
            if format_type == 'bold':
                result_parts.append(f'<b>{match.group(1)}</b>')
            elif format_type == 'italic':
                result_parts.append(f'<i>{match.group(1)}</i>')
            elif format_type == 'code':
                result_parts.append(f'<font name="Courier">{match.group(1)}</font>')
            elif format_type == 'link':
                link_text = match.group(1)
                link_url = match.group(2)
                result_parts.append(f'<link href="{link_url}">{link_text}</link>')
            
            # Move position past the formatted text
            current_pos = pos + len(match.group(0))
        
        return ''.join(result_parts)
    else:
        # For Word output - return plain text without HTML tags
        # Remove markdown formatting but keep the text
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove bold
        text = re.sub(r'\*(.*?)\*', r'\1', text)      # Remove italic
        text = re.sub(r'`(.*?)`', r'\1', text)        # Remove code
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # Remove links, keep text
        return text


def clean_file_structure(text):
    """Clean up file structure diagrams to be more readable."""
    # Replace common file structure patterns
    text = re.sub(r'nnn\s+', '    ', text)  # Replace 'nnn' with proper indentation
    text = re.sub(r'n\s+n\s+n\s+', '        ', text)  # Replace 'n n n' with deeper indentation
    text = re.sub(r'n\s+n\s+', '    ', text)  # Replace 'n n' with indentation
    
    # Clean up any remaining 'n' characters that are likely meant to be spaces
    text = re.sub(r'(?<=\s)n(?=\s)', ' ', text)
    
    # Additional cleaning for file structure patterns
    # Replace patterns like "markdownexporter/ nnn src/" with proper indentation
    text = re.sub(r'(\w+/\s*)nnn\s+', r'\1    ', text)
    text = re.sub(r'(\w+/\s*)n\s+n\s+n\s+', r'\1        ', text)
    text = re.sub(r'(\w+/\s*)n\s+n\s+', r'\1    ', text)
    
    # Clean up any remaining isolated 'n' characters in file paths
    text = re.sub(r'(\w+/\s*)n(?=\s)', r'\1 ', text)
    
    return text


class BaseConverter(ABC):
    """Base class for format converters."""
    
    def __init__(self, output_dir: Path):
        """Initialize converter.
        
        Args:
            output_dir: Output directory for converted files
        """
        self.output_dir = output_dir
        self.output_dir.mkdir(exist_ok=True)
    
    @abstractmethod
    def convert(self, content: str, output_path: Path) -> bool:
        """Convert markdown content to target format.
        
        Args:
            content: Processed markdown content
            output_path: Output file path
            
        Returns:
            True if conversion successful, False otherwise
        """
        pass


class WordConverter(BaseConverter):
    """Word document converter (Primary Focus)."""
    
    def convert(self, content: str, output_path: Path) -> bool:
        """Convert markdown content to Word document.
        
        Args:
            content: Processed markdown content
            output_path: Output file path
            
        Returns:
            True if conversion successful, False otherwise
        """
        if not CONVERSION_AVAILABLE:
            logger.error("Conversion libraries not available")
            return False
            
        try:
            # Create Word document
            doc = Document()
            
            # Add header and footer if configuration is available
            if HEADER_FOOTER_AVAILABLE:
                self._add_header_footer(doc)
            
            # Clean file structure diagrams first
            content = clean_file_structure(content)
            
            # Process Mermaid blocks first
            mermaid_pattern = r'```mermaid\s*\n(.*?)\n```'
            
            def replace_mermaid(match):
                mermaid_code = match.group(1)
                logger.info(f"Rendering Mermaid diagram ({len(mermaid_code)} characters)...")
                
                # Create temporary directory for images
                images_dir = "temp_images"
                Path(images_dir).mkdir(exist_ok=True)
                
                # Render the diagram
                image_path = render_mermaid_diagram(mermaid_code, images_dir)
                
                if image_path:
                    # Add the image to the document
                    try:
                        doc.add_picture(image_path, width=Pt(400))
                        doc.add_paragraph()  # Add space after image
                        return ""  # Remove the original mermaid block
                    except Exception as img_error:
                        logger.error(f"Error adding image to Word: {img_error}")
                        return f"\n**[Mermaid Diagram]**\n\n*Image could not be added to Word.*\n\n"
                else:
                    return f"\n**[Mermaid Diagram]**\n\n*Failed to render diagram.*\n\n"
            
            # Replace Mermaid blocks with rendered images
            content = re.sub(mermaid_pattern, replace_mermaid, content, flags=re.DOTALL)
            
            # Split content into lines
            lines = content.split('\n')
            
            for line in lines:
                line = line.rstrip()
                
                # Handle headers
                if line.startswith('# '):
                    formatted_text = process_formatted_text(line[2:], "word")
                    heading = doc.add_heading(formatted_text, level=1)
                    heading.style.font.size = Pt(16)
                elif line.startswith('## '):
                    formatted_text = process_formatted_text(line[3:], "word")
                    heading = doc.add_heading(formatted_text, level=2)
                    heading.style.font.size = Pt(14)
                elif line.startswith('### '):
                    formatted_text = process_formatted_text(line[4:], "word")
                    heading = doc.add_heading(formatted_text, level=3)
                    heading.style.font.size = Pt(12)
                
                # Handle lists
                elif line.startswith('- ') or line.startswith('* '):
                    formatted_text = process_formatted_text(line[2:], "word")
                    doc.add_paragraph(formatted_text, style='List Bullet')
                elif re.match(r'^\d+\. ', line):
                    match = re.match(r'^(\d+)\. (.*)', line)
                    if match:
                        number, text = match.groups()
                        formatted_text = process_formatted_text(text, "word")
                        doc.add_paragraph(formatted_text, style='List Number')
                
                # Handle empty lines
                elif not line.strip():
                    doc.add_paragraph()
                
                # Handle regular paragraphs
                else:
                    formatted_text = process_formatted_text(line, "word")
                    if formatted_text.strip():
                        doc.add_paragraph(formatted_text)
            
            # Save the document
            doc.save(str(output_path))
            
            # Clean up temporary images
            import shutil
            if os.path.exists("temp_images"):
                shutil.rmtree("temp_images")
            
            return True
            
        except Exception as e:
            logger.error(f"Error converting markdown to Word: {e}")
            return False
    
    def _add_header_footer(self, doc: Document) -> None:
        """Add header and footer to Word document.
        
        Args:
            doc: Word document to add header/footer to
        """
        try:
            config = HeaderFooterConfig()
            
            # Add header
            header_config = config.get_header_config()
            header_text = config.get_header_text()
            
            # Create header paragraph
            header_paragraph = doc.sections[0].header.paragraphs[0]
            header_paragraph.text = header_text
            
            # Apply header formatting
            header_run = header_paragraph.runs[0]
            header_run.font.name = header_config.get("font_name", "Arial")
            header_run.font.size = Pt(header_config.get("font_size", 10))
            header_run.font.bold = header_config.get("bold", False)
            header_run.font.italic = header_config.get("italic", False)
            
            # Set alignment
            alignment = header_config.get("alignment", "center")
            if alignment == "center":
                header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignment == "right":
                header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add footer
            footer_config = config.get_footer_config()
            footer_text = config.get_footer_text()
            
            # Create footer paragraph
            footer_paragraph = doc.sections[0].footer.paragraphs[0]
            footer_paragraph.text = footer_text
            
            # Apply footer formatting
            footer_run = footer_paragraph.runs[0]
            footer_run.font.name = footer_config.get("font_name", "Arial")
            footer_run.font.size = Pt(footer_config.get("font_size", 9))
            footer_run.font.bold = footer_config.get("bold", False)
            footer_run.font.italic = footer_config.get("italic", False)
            
            # Set alignment
            alignment = footer_config.get("alignment", "center")
            if alignment == "center":
                footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignment == "right":
                footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            logger.info("✅ Added header and footer to Word document")
            
        except Exception as e:
            logger.error(f"Error adding header/footer to Word: {e}")
    
    def _parse_table(self, lines: list, start_index: int) -> Optional[list]:
        """Parse table from markdown lines.
        
        Args:
            lines: All lines of content
            start_index: Starting index for table
            
        Returns:
            Table data as list of rows, or None if not a valid table
        """
        table_data = []
        i = start_index
        
        # Check if we have a table
        while i < len(lines) and lines[i].strip().startswith('|'):
            row = lines[i].strip()
            if not row.endswith('|'):
                break
            
            # Parse row
            cells = [cell.strip() for cell in row.split('|')[1:-1]]
            table_data.append(cells)
            i += 1
            
            # Check for separator line
            if i < len(lines) and lines[i].strip().startswith('|'):
                separator = lines[i].strip()
                if all(cell.strip().replace('-', '').replace(':', '') == '' 
                       for cell in separator.split('|')[1:-1]):
                    i += 1
                    continue
                else:
                    break
        
        return table_data if len(table_data) > 1 else None
    
    def _add_table(self, doc: Document, table_data: list) -> None:
        """Add table to Word document.
        
        Args:
            doc: Word document
            table_data: Table data as list of rows
        """
        if not table_data:
            return
        
        rows = len(table_data)
        cols = len(table_data[0])
        
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                cell = row.cells[j]
                cell.text = cell_data
                # Set left alignment for all cells
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _parse_code_block(self, lines: list, start_index: int) -> Optional[list]:
        """Parse code block from markdown lines.
        
        Args:
            lines: All lines of content
            start_index: Starting index for code block
            
        Returns:
            Code content as list of lines, or None if not a valid code block
        """
        if start_index >= len(lines) or not lines[start_index].startswith('```'):
            return None
        
        code_lines = []
        i = start_index + 1
        
        while i < len(lines):
            if lines[i].startswith('```'):
                break
            code_lines.append(lines[i])
            i += 1
        
        return code_lines
    
    def _add_code_block(self, doc: Document, code_lines: list) -> None:
        """Add code block to Word document.
        
        Args:
            doc: Word document
            code_lines: Code content as list of lines
        """
        code_paragraph = doc.add_paragraph()
        code_paragraph.style = 'No Spacing'
        
        for line in code_lines:
            run = code_paragraph.add_run(line + '\n')
            run.font.name = 'Courier New'
            run.font.size = docx.shared.Pt(10)
    
    def _add_image(self, doc: Document, image_line: str) -> None:
        """Add image to Word document.
        
        Args:
            doc: Word document
            image_line: Markdown image line
        """
        # Parse image line: ![alt](url)
        import re
        match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', image_line)
        if match:
            alt_text = match.group(1)
            image_path = match.group(2)
            
            # Check if image file exists
            if Path(image_path).exists():
                try:
                    doc.add_picture(image_path, width=Inches(6))
                    # Add caption
                    caption = doc.add_paragraph(alt_text)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    print(f"Error adding image {image_path}: {e}")


class PDFConverter(BaseConverter):
    """PDF converter with full markdown support."""
    
    def convert(self, content: str, output_path: Path) -> bool:
        """Convert markdown content to PDF document with proper formatting.
        
        Args:
            content: Processed markdown content
            output_path: Output file path
            
        Returns:
            True if conversion successful, False otherwise
        """
        if not CONVERSION_AVAILABLE:
            logger.error("Conversion libraries not available")
            return False
        
        try:
            # Create PDF document with header/footer
            if HEADER_FOOTER_AVAILABLE:
                doc = self._create_pdf_with_header_footer(str(output_path))
            else:
                doc = SimpleDocTemplate(str(output_path), pagesize=A4)
            
            story = []
            
            # Get default styles
            styles = getSampleStyleSheet()
            normal_style = styles['Normal']
            
            # Create custom styles
            code_style = ParagraphStyle(
                'CodeStyle',
                parent=normal_style,
                fontName='Courier',
                fontSize=9,
                backColor=colors.lightgrey,
                leftIndent=20,
                rightIndent=20
            )
            
            file_structure_style = ParagraphStyle(
                'FileStructure',
                parent=normal_style,
                fontName='Courier',
                fontSize=9,
                leftIndent=20
            )
            
            # Clean file structure diagrams first
            content = clean_file_structure(content)
            
            # Process Mermaid blocks first
            mermaid_pattern = r'```mermaid\s*\n(.*?)\n```'
            
            def replace_mermaid(match):
                mermaid_code = match.group(1)
                logger.info(f"Rendering Mermaid diagram ({len(mermaid_code)} characters)...")
                
                # Create temporary directory for images
                images_dir = "temp_images"
                Path(images_dir).mkdir(exist_ok=True)
                
                # Render the diagram
                image_path = render_mermaid_diagram(mermaid_code, images_dir)
                
                if image_path:
                    # Add the image to the document
                    try:
                        img = Image(image_path, width=6*inch, height=4*inch)
                        story.append(img)
                        story.append(Spacer(1, 12))
                        return ""  # Remove the original mermaid block
                    except Exception as img_error:
                        logger.error(f"Error adding image to PDF: {img_error}")
                        return f"\n**[Mermaid Diagram]**\n\n*Image could not be added to PDF.*\n\n"
                else:
                    return f"\n**[Mermaid Diagram]**\n\n*Failed to render diagram.*\n\n"
            
            # Replace Mermaid blocks with rendered images
            content = re.sub(mermaid_pattern, replace_mermaid, content, flags=re.DOTALL)
            
            # Split content into lines
            lines = content.split('\n')
            
            in_code_block = False
            code_lines = []
            
            for line in lines:
                line = line.rstrip()
                
                # Handle code blocks
                if line.startswith('```'):
                    if in_code_block:
                        # End code block
                        if code_lines:
                            code_text = '\n'.join(code_lines)
                            story.append(Paragraph(f'<font name="Courier">{code_text}</font>', code_style))
                            story.append(Spacer(1, 12))
                            code_lines = []
                        in_code_block = False
                    else:
                        # Start code block
                        in_code_block = True
                    continue
                
                if in_code_block:
                    code_lines.append(line)
                    continue
                
                # Handle headers
                if line.startswith('# '):
                    formatted_text = process_formatted_text(line[2:])
                    story.append(Paragraph(formatted_text, styles['Heading1']))
                    story.append(Spacer(1, 12))
                elif line.startswith('## '):
                    formatted_text = process_formatted_text(line[3:])
                    story.append(Paragraph(formatted_text, styles['Heading2']))
                    story.append(Spacer(1, 12))
                elif line.startswith('### '):
                    formatted_text = process_formatted_text(line[4:])
                    story.append(Paragraph(formatted_text, styles['Heading3']))
                    story.append(Spacer(1, 12))
                elif line.startswith('#### '):
                    formatted_text = process_formatted_text(line[5:])
                    story.append(Paragraph(formatted_text, styles['Heading4']))
                    story.append(Spacer(1, 12))
                
                # Handle lists
                elif line.startswith('- ') or line.startswith('* '):
                    formatted_text = process_formatted_text(line[2:])
                    bullet_style = ParagraphStyle('Bullet', parent=normal_style, leftIndent=20)
                    story.append(Paragraph(f'• {formatted_text}', bullet_style))
                elif re.match(r'^\d+\. ', line):
                    match = re.match(r'^(\d+)\. (.*)', line)
                    if match:
                        number, text = match.groups()
                        formatted_text = process_formatted_text(text)
                        bullet_style = ParagraphStyle('Number', parent=normal_style, leftIndent=20)
                        story.append(Paragraph(f'{number}. {formatted_text}', bullet_style))
                
                # Handle file structure lines (lines with indentation and file/directory names)
                elif re.match(r'^\s+[a-zA-Z0-9_\-\./]+', line):
                    # This looks like a file structure line
                    story.append(Paragraph(line, file_structure_style))
                
                # Handle empty lines
                elif not line.strip():
                    story.append(Spacer(1, 6))
                
                # Handle regular paragraphs
                else:
                    formatted_text = process_formatted_text(line)
                    if formatted_text.strip():
                        story.append(Paragraph(formatted_text, normal_style))
                        story.append(Spacer(1, 6))
            
            # Build the PDF
            doc.build(story)
            
            # Clean up temporary images
            import shutil
            if os.path.exists("temp_images"):
                shutil.rmtree("temp_images")
            
            return True
            
        except Exception as e:
            logger.error(f"Error converting markdown to PDF: {e}")
        return False
    
    def _create_pdf_with_header_footer(self, output_path: str) -> SimpleDocTemplate:
        """Create PDF document with header and footer.
        
        Args:
            output_path: Output file path
            
        Returns:
            SimpleDocTemplate with header/footer
        """
        try:
            config = HeaderFooterConfig()
            page_config = config.get_page_config()
            
            # Create custom page template with header/footer
            from reportlab.platypus import PageTemplate, Frame
            from reportlab.lib.units import inch
            
            # Define page margins
            margin_top = page_config.get("margin_top", 1.0) * inch
            margin_bottom = page_config.get("margin_bottom", 1.0) * inch
            margin_left = page_config.get("margin_left", 1.0) * inch
            margin_right = page_config.get("margin_right", 1.0) * inch
            
            # Create frame for content
            frame = Frame(
                margin_left, margin_bottom,
                A4[0] - margin_left - margin_right,
                A4[1] - margin_top - margin_bottom,
                id='normal'
            )
            
            # Create page template with header/footer
            def header_footer(canvas, doc):
                # Add header
                header_config = config.get_header_config()
                header_text = config.get_header_text()
                
                canvas.saveState()
                canvas.setFont(header_config.get("font_name", "Helvetica"), 
                              header_config.get("font_size", 10))
                canvas.setFillColor(header_config.get("color", "#333333"))
                
                # Position header
                header_y = A4[1] - margin_top + 0.2 * inch
                canvas.drawCentredString(A4[0]/2, header_y, header_text)
                
                # Add footer
                footer_config = config.get_footer_config()
                footer_text = config.get_footer_text()
                
                canvas.setFont(footer_config.get("font_name", "Helvetica"), 
                              footer_config.get("font_size", 9))
                canvas.setFillColor(footer_config.get("color", "#666666"))
                
                # Position footer
                footer_y = margin_bottom - 0.2 * inch
                canvas.drawCentredString(A4[0]/2, footer_y, footer_text)
                
                canvas.restoreState()
            
            # Create page template
            page_template = PageTemplate(
                id='header_footer',
                frames=[frame],
                onPage=header_footer
            )
            
            # Create document with template
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            doc.addPageTemplates([page_template])
            
            logger.info("✅ Created PDF document with header and footer")
            return doc
            
        except Exception as e:
            logger.error(f"Error creating PDF with header/footer: {e}")
            # Fallback to simple document
            return SimpleDocTemplate(output_path, pagesize=A4)


class HTMLConverter(BaseConverter):
    """HTML converter for markdown content."""
    
    def convert(self, content: str, output_path: Path) -> bool:
        """Convert markdown content to HTML.
        
        Args:
            content: Processed markdown content
            output_path: Output file path
            
        Returns:
            True if conversion successful, False otherwise
        """
        try:
            # Process the content first
            processed_content = self._process_content(content)
            
            # Create HTML document
            html_content = self._create_html_document(processed_content)
            
            # Write to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            logger.info(f"✅ Successfully converted markdown to HTML: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error converting markdown to HTML: {e}")
        return False
    
    def _process_content(self, content: str) -> str:
        """Process markdown content for HTML conversion."""
        # Process Mermaid diagrams FIRST (before any other text processing)
        content = self._process_mermaid_diagrams(content)
        
        # Clean file structure
        content = clean_file_structure(content)
        
        # Process formatted text
        content = process_formatted_text(content)
        
        return content
    
    def _process_mermaid_diagrams(self, content: str) -> str:
        """Process Mermaid diagrams in the content."""
        import re
        
        def replace_mermaid_block(match):
            mermaid_code = match.group(1)
            logger.info(f"Rendering Mermaid diagram ({len(mermaid_code)} characters)...")
            
            try:
                # Create temp_images directory if it doesn't exist
                images_dir = "temp_images"
                os.makedirs(images_dir, exist_ok=True)
                
                # Render the diagram
                image_path = render_mermaid_diagram(mermaid_code, images_dir)
                
                if image_path and os.path.exists(image_path):
                    # Copy image to results/images directory for HTML access
                    import shutil
                    import uuid
                    results_image_name = f"diagram_{uuid.uuid4().hex[:8]}.png"
                    images_dir = os.path.join(str(self.output_dir), "images")
                    os.makedirs(images_dir, exist_ok=True)
                    results_image_path = os.path.join(images_dir, results_image_name)
                    shutil.copy2(image_path, results_image_path)
                    
                    # Return HTML img tag with relative path to images folder
                    return f'<img src="images/{results_image_name}" alt="Mermaid Diagram" style="max-width: 100%; height: auto;">'
                else:
                    # If rendering failed, return a placeholder with the original code
                    return f'<div style="background-color: #f5f5f5; padding: 15px; border: 1px solid #ddd; border-radius: 5px; margin: 10px 0;"><p><strong>Mermaid Diagram (Rendering Failed):</strong></p><pre><code>{mermaid_code}</code></pre></div>'
            except Exception as e:
                logger.error(f"Error processing Mermaid diagram: {e}")
                # Return a placeholder with the original code
                return f'<div style="background-color: #f5f5f5; padding: 15px; border: 1px solid #ddd; border-radius: 5px; margin: 10px 0;"><p><strong>Mermaid Diagram (Error):</strong></p><pre><code>{mermaid_code}</code></pre></div>'
        
        # Find and replace Mermaid code blocks
        pattern = r'```mermaid\s*\n(.*?)\n```'
        return re.sub(pattern, replace_mermaid_block, content, flags=re.DOTALL)
    
    def _create_html_document(self, content: str) -> str:
        """Create a complete HTML document from processed content."""
        # Split content into lines
        lines = content.split('\n')
        html_lines = []
        
        # HTML header
        html_lines.append('<!DOCTYPE html>')
        html_lines.append('<html lang="en">')
        html_lines.append('<head>')
        html_lines.append('    <meta charset="UTF-8">')
        html_lines.append('    <meta name="viewport" content="width=device-width, initial-scale=1.0">')
        html_lines.append('    <title>Markdown Export</title>')
        html_lines.append('    <style>')
        html_lines.append('        body { font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; }')
        html_lines.append('        h1, h2, h3, h4 { color: #333; }')
        html_lines.append('        h1 { border-bottom: 2px solid #eee; padding-bottom: 10px; }')
        html_lines.append('        h2 { border-bottom: 1px solid #eee; padding-bottom: 5px; }')
        html_lines.append('        code { background-color: #f4f4f4; padding: 2px 4px; border-radius: 3px; font-family: "Courier New", monospace; }')
        html_lines.append('        pre { background-color: #f4f4f4; padding: 15px; border-radius: 5px; overflow-x: auto; }')
        html_lines.append('        pre code { background-color: transparent; padding: 0; }')
        html_lines.append('        ul, ol { padding-left: 20px; }')
        html_lines.append('        li { margin: 5px 0; }')
        html_lines.append('        a { color: #007acc; text-decoration: none; }')
        html_lines.append('        a:hover { text-decoration: underline; }')
        html_lines.append('        .file-structure { font-family: "Courier New", monospace; background-color: #f9f9f9; padding: 10px; border-radius: 5px; }')
        html_lines.append('        img { max-width: 100%; height: auto; }')
        html_lines.append('    </style>')
        html_lines.append('</head>')
        html_lines.append('<body>')
        
        in_code_block = False
        code_lines = []
        
        for line in lines:
            line = line.rstrip()
            
            # Handle code blocks
            if line.startswith('```'):
                if in_code_block:
                    # End code block
                    if code_lines:
                        html_lines.append('<pre><code>')
                        html_lines.extend(code_lines)
                        html_lines.append('</code></pre>')
                        code_lines = []
                    in_code_block = False
                else:
                    # Start code block
                    in_code_block = True
                continue
            
            if in_code_block:
                code_lines.append(line)
                continue
            
            # Handle headers
            if line.startswith('# '):
                html_lines.append(f'<h1>{line[2:]}</h1>')
            elif line.startswith('## '):
                html_lines.append(f'<h2>{line[3:]}</h2>')
            elif line.startswith('### '):
                html_lines.append(f'<h3>{line[4:]}</h3>')
            elif line.startswith('#### '):
                html_lines.append(f'<h4>{line[5:]}</h4>')
            
            # Handle lists
            elif line.startswith('- ') or line.startswith('* '):
                html_lines.append(f'<li>{line[2:]}</li>')
            elif re.match(r'^\d+\. ', line):
                match = re.match(r'^(\d+)\. (.*)', line)
                if match:
                    number, text = match.groups()
                    html_lines.append(f'<li>{text}</li>')
            
            # Handle file structure lines
            elif re.match(r'^\s+[a-zA-Z0-9_\-\./]+', line):
                html_lines.append(f'<div class="file-structure">{line}</div>')
            
            # Handle empty lines
            elif not line.strip():
                html_lines.append('<br>')
            
            # Handle regular paragraphs
            else:
                if line.strip():
                    html_lines.append(f'<p>{line}</p>')
        
        # Close HTML
        html_lines.append('</body>')
        html_lines.append('</html>')
        
        return '\n'.join(html_lines)
